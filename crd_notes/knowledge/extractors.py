from __future__ import annotations

import csv
import re
from dataclasses import dataclass, field
from pathlib import Path

from crd_notes.core.errors import ConfigurationError

MAX_EXTRACTED_CHARS = 1_000_000
MAX_PDF_PAGES = 250
MAX_CSV_ROWS = 20_000
MAX_XLSX_ROWS = 20_000

SUPPORTED_KNOWLEDGE_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".txt",
    ".md",
    ".xls",
    ".xlsx",
    ".csv",
}


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    source_type: str
    page_count: int | None = None
    sheet_count: int | None = None
    row_count: int | None = None
    warnings: list[str] = field(default_factory=list)
    metadata: dict[str, str | int | float | bool] = field(default_factory=dict)


def extract_text_from_file(path: Path) -> str:
    return extract_document_from_file(path).text


def extract_document_from_file(path: Path) -> ExtractedDocument:
    extension = path.suffix.lower()
    if extension == ".doc":
        raise ConfigurationError(
            "Formato Word legacy .doc non supportato in modo affidabile.",
            detail=(
                "Converti il file in .docx, .pdf, .md o .txt prima di importarlo "
                "nella knowledge base."
            ),
        )
    if extension not in SUPPORTED_KNOWLEDGE_EXTENSIONS:
        raise ConfigurationError(
            f"Formato knowledge base non supportato: {extension or 'sconosciuto'}"
        )
    if extension == ".pdf":
        return _extract_pdf(path)
    if extension == ".docx":
        return _extract_docx(path)
    if extension in {".xls", ".xlsx"}:
        return _extract_xlsx(path)
    if extension == ".csv":
        return _extract_csv(path)
    source_type = "markdown" if extension == ".md" else "text"
    return _extract_text(path, source_type=source_type)


def _extract_pdf(path: Path) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ConfigurationError(
            "Dipendenza mancante per PDF: pypdf.",
            detail="Installa le dipendenze aggiornate del progetto.",
        ) from exc
    reader = PdfReader(str(path))
    pages: list[str] = []
    warnings: list[str] = []
    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        warnings.append(f"PDF tagliato alle prime {MAX_PDF_PAGES} pagine.")
    for index, page in enumerate(reader.pages, start=1):
        if index > MAX_PDF_PAGES:
            break
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Pagina {index}]\n{text}")
    text, truncated = _limit_text("\n\n".join(pages))
    if truncated:
        warnings.append(f"Testo estratto tagliato a {MAX_EXTRACTED_CHARS} caratteri.")
    return ExtractedDocument(
        text=text,
        source_type="pdf",
        page_count=page_count,
        warnings=warnings,
    )


def _extract_docx(path: Path) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise ConfigurationError(
            "Dipendenza mancante per Word: python-docx.",
            detail="Installa le dipendenze aggiornate del progetto.",
        ) from exc
    try:
        document = Document(str(path))
    except Exception as exc:
        raise ConfigurationError(
            "Documento Word non leggibile.",
            detail="Il file non sembra un .docx valido o e' danneggiato.",
        ) from exc
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                parts.append(" | ".join(values))
    text, truncated = _limit_text("\n".join(parts))
    warnings = [f"Testo estratto tagliato a {MAX_EXTRACTED_CHARS} caratteri."] if truncated else []
    return ExtractedDocument(
        text=text,
        source_type="docx",
        warnings=warnings,
        metadata={
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
    )


def _extract_xlsx(path: Path) -> ExtractedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise ConfigurationError(
            "Dipendenza mancante per Excel: openpyxl.",
            detail="Installa le dipendenze aggiornate del progetto.",
        ) from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[str] = []
    warnings: list[str] = []
    row_count = 0
    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            if row_count >= MAX_XLSX_ROWS:
                warnings.append(f"Excel tagliato alle prime {MAX_XLSX_ROWS} righe.")
                break
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))
                row_count += 1
        if rows:
            blocks.append(f"[Sheet: {sheet.title}]\n" + "\n".join(rows))
        if row_count >= MAX_XLSX_ROWS:
            break
    text, truncated = _limit_text("\n\n".join(blocks))
    if truncated:
        warnings.append(f"Testo estratto tagliato a {MAX_EXTRACTED_CHARS} caratteri.")
    return ExtractedDocument(
        text=text,
        source_type="xlsx",
        sheet_count=len(workbook.worksheets),
        row_count=row_count,
        warnings=warnings,
    )


def _extract_csv(path: Path) -> ExtractedDocument:
    decode_errors: list[str] = []
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            with path.open("r", encoding=encoding, newline="") as file:
                reader = csv.reader(file)
                rows: list[str] = []
                truncated_rows = False
                for index, row in enumerate(reader, start=1):
                    if index > MAX_CSV_ROWS:
                        truncated_rows = True
                        break
                    value = " | ".join(cell.strip() for cell in row if cell and cell.strip())
                    if value:
                        rows.append(value)
            lines = [row for row in rows if row]
            if lines:
                warnings = []
                if truncated_rows:
                    warnings.append(f"CSV tagliato alle prime {MAX_CSV_ROWS} righe.")
                text, truncated = _limit_text("\n".join(lines))
                if truncated:
                    warnings.append(
                        f"Testo estratto tagliato a {MAX_EXTRACTED_CHARS} caratteri."
                    )
                return ExtractedDocument(
                    text=text,
                    source_type="csv",
                    row_count=len(lines),
                    warnings=warnings,
                    metadata={"encoding": encoding},
                )
        except UnicodeDecodeError as exc:
            decode_errors.append(f"{encoding}: {exc}")
            continue
    raise ConfigurationError(
        "CSV non leggibile.",
        detail="; ".join(decode_errors) or "Nessuna riga testuale estratta dal CSV.",
    )


def _extract_text(path: Path, *, source_type: str = "text") -> ExtractedDocument:
    decode_errors: list[str] = []
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            text, truncated = _limit_text(path.read_text(encoding=encoding))
            warnings = (
                [f"Testo estratto tagliato a {MAX_EXTRACTED_CHARS} caratteri."]
                if truncated
                else []
            )
            return ExtractedDocument(
                text=text,
                source_type=source_type,
                warnings=warnings,
                metadata={"encoding": encoding},
            )
        except UnicodeDecodeError as exc:
            decode_errors.append(f"{encoding}: {exc}")
            continue
    raise ConfigurationError(
        "File testuale non leggibile.",
        detail="; ".join(decode_errors) or "Encoding non riconosciuto.",
    )


def _limit_text(value: str) -> tuple[str, bool]:
    cleaned = re.sub(r"\s+\n", "\n", value).strip()
    if len(cleaned) <= MAX_EXTRACTED_CHARS:
        return cleaned, False
    return cleaned[:MAX_EXTRACTED_CHARS].rstrip(), True
